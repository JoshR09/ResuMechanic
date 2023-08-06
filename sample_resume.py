from docx import Document
from docx.shared import Cm
from docx.shared import Pt
from docx.oxml import parse_xml


def set_heading_style(paragraph):
    run = paragraph.runs[0]
    font = run.font
    font.name = 'Garamond'
    font.size = Pt(22)
    font.bold = True


def set_section_heading_style(paragraph):
    paragraph.paragraph_format.space_after = Pt(3)  # Remove space after paragraph
    paragraph.paragraph_format.space_before = Pt(3)  # Remove space before paragraph
    run = paragraph.runs[0]
    font = run.font
    font.name = 'Garamond'
    font.size = Pt(11.5)
    font.bold = True


def set_paragraph_style(paragraph):
    paragraph.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
    paragraph.paragraph_format.space_before = Pt(0)  # Remove space before paragraph
    for run in paragraph.runs:
        font = run.font
        font.name = 'Garamond'
        font.size = Pt(11.5)


def set_bold_paragraph_style(paragraph):
    paragraph.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
    paragraph.paragraph_format.space_before = Pt(0)  # Remove space before paragraph
    for run in paragraph.runs:
        font = run.font
        font.name = 'Garamond'
        font.size = Pt(12)
        font.bold = True


# Function to add a horizontal line after the section heading
def add_horizontal_line(paragraph):
    p = paragraph._element
    bottom_border = parse_xml(
        f'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/>'
        f'</w:pBdr>'
    )
    pPr = p.get_or_add_pPr()
    pPr.append(bottom_border)


def test_resume_generator():
    # Set up test data
    name = "John Doe"
    phone = "1234567890"
    email = "johndoe@example.com"
    postcode = "12345"
    summary = "Motivated and enthusiastic student eager to gain experience in an active and rewarding environment. Able to communicate information clearly and possessing strong time management abilities, as well as adaptability to new skills and concepts."
    employment_history = [("Company A", "Software Engineer", "01/2018", "12/2020", ["Developed web applications", "Led a team of developers"]), ("Company B", "Senior Software Engineer", "01/2021", "Present", ["Designed system architecture", "Implemented new features"])]
    educational_history = [("University A", "Bachelor of Science", "05/2017", ["Graduated with honors"]), ("University B", "Master of Engineering", "05/2019", ["Thesis on machine learning"])]
    skills = [("Python", "Proficient in Python programming"), ("Java", "Experience with Java development")]

    # Create a new Word document
    document = Document()

    # Set smaller margins
    sections = document.sections
    for section in sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    # Add title
    set_heading_style(document.add_paragraph(f'{name.upper()}'))

    # Add contact information section heading
    section_heading = document.add_paragraph('CONTACT INFORMATION')
    set_section_heading_style(section_heading)
    add_horizontal_line(section_heading)  # Add horizontal line after the section heading

    # Add contact information details
    contact_info_paragraphs = [
        document.add_paragraph(f'Phone: {phone}'),
        document.add_paragraph(f'Email: {email}'),
        document.add_paragraph(f'Address: {postcode}'),
        document.add_paragraph(f'')
    ]
    for paragraph in contact_info_paragraphs:
        set_paragraph_style(paragraph)

    # Add personal summary section heading
    section_heading = document.add_paragraph('PERSONAL SUMMARY')
    set_section_heading_style(section_heading)
    add_horizontal_line(section_heading)  # Add horizontal line after the section heading

    # Add personal summary
    summary_paragraph = document.add_paragraph(summary)
    ps_whitespace = document.add_paragraph(f'')
    set_paragraph_style(summary_paragraph)
    set_paragraph_style(ps_whitespace)

    # Add employment history section heading
    section_heading = document.add_paragraph('EMPLOYMENT HISTORY')
    set_section_heading_style(section_heading)
    add_horizontal_line(section_heading)  # Add horizontal line after the section heading

    # Add employment history
    for employer, title, start_date, end_date, responsibilities in employment_history:
        employment_info_paragraph = document.add_paragraph(f'')
        title_run = employment_info_paragraph.add_run(title)
        title_run.bold = True
        employment_info_paragraph.add_run(f' | {employer}, {start_date} - {end_date}')
        for responsibility in responsibilities:
            responsibilities_paragraph = document.add_paragraph(responsibility)
            set_paragraph_style(responsibilities_paragraph)
        set_paragraph_style(employment_info_paragraph)
        m_whitespace = document.add_paragraph(f'')
        set_paragraph_style(m_whitespace)

    # Add educational history section heading
    section_heading = document.add_paragraph('EDUCATIONAL HISTORY')
    set_section_heading_style(section_heading)
    add_horizontal_line(section_heading)  # Add horizontal line after the section heading

    # Add educational history
    for institution, degree, completion_date, achievements in educational_history:
        educational_info_paragraph = document.add_paragraph(f'')
        institution_run = educational_info_paragraph.add_run(institution)
        institution_run.bold = True
        educational_info_paragraph.add_run(f' | {degree}, {completion_date}')
        for achievement in achievements:
            achievements_paragraph = document.add_paragraph(achievement)
            set_paragraph_style(achievements_paragraph)
        set_paragraph_style(educational_info_paragraph)
        e_whitespace = document.add_paragraph(f'')
        set_paragraph_style(e_whitespace)

    # Add skills section heading
    section_heading = document.add_paragraph('SKILLS')
    set_section_heading_style(section_heading)
    add_horizontal_line(section_heading)  # Add horizontal line after the section heading

    # Add skills
    for skill, description in skills:
        skill_paragraph = document.add_paragraph(skill)
        set_bold_paragraph_style(skill_paragraph)
        description_paragraph = document.add_paragraph(description)
        set_paragraph_style(description_paragraph)
        s_whitespace = document.add_paragraph(f'')
        set_paragraph_style(s_whitespace)

    # Save the document
    document.save('resume_test.docx')
    print("Resume created successfully!")


# Call the function to test the resume generator
test_resume_generator()
