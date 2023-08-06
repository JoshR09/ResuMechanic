from docx import Document
from docx.shared import Cm
from docx.shared import Pt
from docx.oxml import parse_xml


def set_heading_style(paragraph):
    run = paragraph.runs[0]
    font = run.font
    font.name = 'Garamond'
    font.size = Pt(22)
    font.bold = True


def set_section_heading_style(paragraph):
    paragraph.paragraph_format.space_after = Pt(3)  # Remove space after paragraph
    paragraph.paragraph_format.space_before = Pt(3)  # Remove space before paragraph
    run = paragraph.runs[0]
    font = run.font
    font.name = 'Garamond'
    font.size = Pt(12)
    font.bold = True


def set_paragraph_style(paragraph):
    paragraph.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
    paragraph.paragraph_format.space_before = Pt(0)  # Remove space before paragraph
    for run in paragraph.runs:
        font = run.font
        font.name = 'Garamond'
        font.size = Pt(12)


def set_bold_paragraph_style(paragraph):
    paragraph.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
    paragraph.paragraph_format.space_before = Pt(0)  # Remove space before paragraph
    for run in paragraph.runs:
        font = run.font
        font.name = 'Garamond'
        font.size = Pt(12)
        font.bold = True


# Function to add a horizontal line after the section heading
def add_horizontal_line(paragraph):
    p = paragraph._element
    bottom_border = parse_xml(
        f'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/>'
        f'</w:pBdr>'
    )
    pPr = p.get_or_add_pPr()
    pPr.append(bottom_border)


def create_resume():
    # Prompt the user for contact information
    name = input("Enter your full name: ")
    phone = input("Enter your phone number: ")
    email = input("Enter your email address: ")
    postcode = input("Enter your postcode: ")

    # Prompt the user for personal summary
    summary = input("Enter a personal summary: ")

    # Prompt the user for employment history
    print("Enter your employment history")
    employment_history = []
    while True:
        employer = input("Enter employer name (Enter 'q' to stop): ")
        if employer == 'q':
            break
        title = input("Enter job title: ")
        start_date = input("Enter start date (MM/YYYY): ")
        end_date = input("Enter end date (MM/YYYY): ")
        print("Enter your job responsibilities")
        responsibilities = []
        while True:
            responsibility = input("Enter job responsibilities (Enter 'q' to stop): ")
            if responsibility == 'q':
                break
            responsibilities.append(responsibility)
        employment_history.append((employer, title, start_date, end_date, responsibilities))

    # Prompt the user for educational history
    print("Enter your educational history")
    educational_history = []
    while True:
        institution = input("Enter institution name (Enter 'q' to stop): ")
        if institution == 'q':
            break
        degree = input("Enter certification/degree obtained: ")
        completion_date = input("Enter completion date (MM/YYYY): ")
        achievements = []
        while True:
            achievement = input("Enter achievements (Enter 'q' to stop): ")
            if achievement == 'q':
                break
            achievements.append(achievement)
        educational_history.append((institution, degree, completion_date, achievements))

    # Prompt the user for skills
    print("Enter your skills")
    skills = []
    while True:
        skill = input("Enter a skill (Enter 'q' to stop): ")
        if skill == 'q':
            break
        description = input("Enter a description/explanation of the skill:")
        skills.append((skill, description))

    # Create a new Word document
    document = Document()

    # Set smaller margins
    sections = document.sections
    for section in sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    # Add title
    set_heading_style(document.add_paragraph(f'{name.upper()}'))

    # Add contact information section heading
    section_heading = document.add_paragraph('CONTACT INFORMATION')
    set_section_heading_style(section_heading)
    add_horizontal_line(section_heading)  # Add horizontal line after the section heading

    # Add contact information details
    contact_info_paragraphs = [
        document.add_paragraph(f'Phone: {phone}'),
        document.add_paragraph(f'Email: {email}'),
        document.add_paragraph(f'Address: {postcode}'),
        document.add_paragraph(f'')
    ]
    for paragraph in contact_info_paragraphs:
        set_paragraph_style(paragraph)

    # Add personal summary section heading
    section_heading = document.add_paragraph('PERSONAL SUMMARY')
    set_section_heading_style(section_heading)
    add_horizontal_line(section_heading)  # Add horizontal line after the section heading

    # Add personal summary
    summary_paragraph = document.add_paragraph(summary)
    ps_whitespace = document.add_paragraph(f'')
    set_paragraph_style(summary_paragraph)
    set_paragraph_style(ps_whitespace)

    # Add employment history section heading
    section_heading = document.add_paragraph('EMPLOYMENT HISTORY')
    set_section_heading_style(section_heading)
    add_horizontal_line(section_heading)  # Add horizontal line after the section heading

    # Add employment history
    for employer, title, start_date, end_date, responsibilities in employment_history:
        employment_info_paragraph = document.add_paragraph(f'')
        title_run = employment_info_paragraph.add_run(title)
        title_run.bold = True
        employment_info_paragraph.add_run(f' | {employer}, {start_date} - {end_date}')
        for responsibility in responsibilities:
            responsibilities_paragraph = document.add_paragraph(responsibility)
            set_paragraph_style(responsibilities_paragraph)
        set_paragraph_style(employment_info_paragraph)
        m_whitespace = document.add_paragraph(f'')
        set_paragraph_style(m_whitespace)

    # Add educational history section heading
    section_heading = document.add_paragraph('EDUCATIONAL HISTORY')
    set_section_heading_style(section_heading)
    add_horizontal_line(section_heading)  # Add horizontal line after the section heading

    # Add educational history
    for institution, degree, completion_date, achievements in educational_history:
        educational_info_paragraph = document.add_paragraph(f'')
        institution_run = educational_info_paragraph.add_run(institution)
        institution_run.bold = True
        educational_info_paragraph.add_run(f' | {degree}, {completion_date}')
        for achievement in achievements:
            achievements_paragraph = document.add_paragraph(achievement)
            set_paragraph_style(achievements_paragraph)
        set_paragraph_style(educational_info_paragraph)
        e_whitespace = document.add_paragraph(f'')
        set_paragraph_style(e_whitespace)

    # Add skills section heading
    section_heading = document.add_paragraph('SKILLS')
    set_section_heading_style(section_heading)
    add_horizontal_line(section_heading)  # Add horizontal line after the section heading

    # Add skills
    for skill, description in skills:
        skill_paragraph = document.add_paragraph(skill)
        set_bold_paragraph_style(skill_paragraph)
        description_paragraph = document.add_paragraph(description)
        set_paragraph_style(description_paragraph)
        s_whitespace = document.add_paragraph(f'')
        set_paragraph_style(s_whitespace)

    # Save the document
    document.save('resume.docx')
    print("Resume created successfully!")


# Call the function to create the resume
create_resume()
